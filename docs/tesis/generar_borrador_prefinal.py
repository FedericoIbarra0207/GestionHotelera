"""Genera un borrador DOCX técnico para el pre-final de DVT.

El archivo se crea desde contenido propio y datos constatados en el repositorio.
No incorpora credenciales, usuarios reales ni copias de la base de producción.
Abrir el DOCX en Microsoft Word y actualizar los campos (Ctrl+A, F9) para que
la tabla de contenido y los números de página se regeneren antes de entregar.
"""

from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(__file__).resolve().parent
OUTPUT = BASE_DIR / "DVT - Proyecto Final Integrador - Borrador Prefinal 2026.docx"
FIGURES_DIR = BASE_DIR / "figuras"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
GRAY = "F2F2F2"
DARK = "1F2937"
PIL_BLUE = f"#{BLUE}"
PIL_DARK = f"#{DARK}"


def set_cell_shading(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(int(width_cm * 567)))
    width.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:cantSplit")
    tr_pr.append(element)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    field = OxmlElement("w:instrText")
    field.set(qn("xml:space"), "preserve")
    field.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Actualice este campo en Microsoft Word (Ctrl+A, F9)."
    separate.append(placeholder)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, field, separate, end])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def make_font(size=12, bold=False, italic=False, color=None):
    font = {"name": "Times New Roman", "size": Pt(size), "bold": bold, "italic": italic}
    if color:
        font["color"] = RGBColor.from_string(color)
    return font


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.27)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in [("Title", 18), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5
    doc.styles["Heading 1"].paragraph_format.page_break_before = True

    if "Caption" not in [s.name for s in doc.styles]:
        doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title(doc, title, subtitle=None):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.add_run(title)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.paragraph_format.first_line_indent = Cm(0)


def add_text(doc, paragraphs):
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(item)


def add_table(doc, headers, rows, widths=None, caption=None):
    if caption:
        p = doc.add_paragraph(caption, style="Caption")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(headers):
        cell = header.cells[i]
        cell.text = value
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            set_cell_width(cell, widths[i])
        for run in cell.paragraphs[0].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            cell.text = str(value)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if widths:
                set_cell_width(cell, widths[i])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
        if len(table.rows) % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, GRAY)
    doc.add_paragraph("")
    return table


def draw_box(draw, box, text, fill, outline=PIL_BLUE, size=22):
    draw.rounded_rectangle(box, radius=16, fill=fill, outline=outline, width=3)
    font = ImageFont.truetype("arial.ttf", size)
    lines = text.split("\n")
    total_h = sum(draw.textbbox((0, 0), line, font=font)[3] for line in lines) + 8 * (len(lines) - 1)
    y = (box[1] + box[3] - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = (box[0] + box[2] - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill=PIL_DARK, font=font)
        y += (bbox[3] - bbox[1]) + 8


def arrow(draw, start, end):
    draw.line([start, end], fill=PIL_BLUE, width=4)
    x, y = end
    draw.polygon([(x, y), (x - 12, y - 7), (x - 12, y + 7)], fill=PIL_BLUE)


def build_figures():
    FIGURES_DIR.mkdir(exist_ok=True)
    items = {}

    image = Image.new("RGB", (1400, 800), "white")
    draw = ImageDraw.Draw(image)
    draw.text((50, 30), "Arquitectura de solución técnica de DVT", fill=PIL_BLUE, font=ImageFont.truetype("arialbd.ttf", 34))
    boxes = {
        "landing": (60, 180, 320, 340), "spa": (60, 480, 320, 640),
        "api": (530, 300, 850, 500), "db": (1080, 180, 1340, 340), "render": (1080, 480, 1340, 640),
    }
    draw_box(draw, boxes["landing"], "Landing\ncomercial", "#E8F1FA")
    draw_box(draw, boxes["spa"], "Aplicación interna\nVue 3 / Vite", "#E8F1FA")
    draw_box(draw, boxes["api"], "API REST\nNode.js + Express\nJWT / CORS", "#FFF3CD")
    draw_box(draw, boxes["db"], "Cloud Firestore\nFirebase Admin SDK", "#E7F5EC")
    draw_box(draw, boxes["render"], "Render\nWeb Service", "#F8E9F1")
    arrow(draw, (320, 260), (530, 350)); arrow(draw, (320, 560), (530, 450)); arrow(draw, (850, 350), (1080, 260)); arrow(draw, (850, 450), (1080, 560))
    path = FIGURES_DIR / "arquitectura.png"; image.save(path); items["arquitectura"] = path

    image = Image.new("RGB", (1400, 880), "white")
    draw = ImageDraw.Draw(image)
    draw.text((50, 30), "Modelo lógico de datos de Firestore", fill=PIL_BLUE, font=ImageFont.truetype("arialbd.ttf", 34))
    entities = [
        (60, 170, 310, 340, "USUARIOS\nrol · activo"), (420, 120, 700, 300, "HABITACIONES\nnumero · tarifa · estado"),
        (420, 490, 700, 670, "HUESPEDES\ndocumento · contacto"), (820, 300, 1100, 500, "RESERVAS\nfechas · snapshots"),
        (1150, 120, 1370, 300, "PAGOS\nmonto · método"), (1150, 540, 1370, 720, "CONSUMOS_EXTRAS\ncategoría · estado"),
    ]
    for x1, y1, x2, y2, text in entities: draw_box(draw, (x1, y1, x2, y2), text, "#E8F1FA", size=19)
    for start, end in [((700, 210), (820, 350)), ((700, 580), (820, 450)), ((1100, 350), (1150, 250)), ((1100, 450), (1150, 620)), ((310, 250), (820, 390))]: arrow(draw, start, end)
    path = FIGURES_DIR / "modelo-datos.png"; image.save(path); items["datos"] = path

    image = Image.new("RGB", (1400, 700), "white")
    draw = ImageDraw.Draw(image)
    draw.text((50, 30), "Flujo de consumo y pago", fill=PIL_BLUE, font=ImageFont.truetype("arialbd.ttf", 34))
    flow = [(50, 280, 285, 445, "Registro\nPENDIENTE"), (380, 280, 615, 445, "Corrección\noperativa"), (710, 280, 945, 445, "Agregar a cuenta\nEN_CUENTA"), (1040, 280, 1325, 445, "Pago confirmado\nCONFIRMADO")]
    colors = ["#FDECEC", "#FFF3CD", "#E7F5EC", "#E8F1FA"]
    for box, color in zip(flow, colors): draw_box(draw, box[:4], box[4], color, size=18)
    for i in range(3): arrow(draw, (flow[i][2], 360), (flow[i + 1][0], 360))
    path = FIGURES_DIR / "flujo-consumos-pagos.png"; image.save(path); items["consumos"] = path

    image = Image.new("RGB", (1400, 700), "white")
    draw = ImageDraw.Draw(image)
    draw.text((50, 30), "Flujo interno de recuperación de acceso", fill=PIL_BLUE, font=ImageFont.truetype("arialbd.ttf", 34))
    flow = [(45, 280, 300, 445, "Solicitud\ndesde login"), (385, 280, 640, 445, "Registro interno\nPENDIENTE"), (725, 280, 980, 445, "Validación por\nadministrador"), (1065, 280, 1350, 445, "Clave temporal y\ncambio obligatorio")]
    for box, color in zip(flow, colors): draw_box(draw, box[:4], box[4], color, size=18)
    for i in range(3): arrow(draw, (flow[i][2], 360), (flow[i + 1][0], 360))
    path = FIGURES_DIR / "flujo-recuperacion.png"; image.save(path); items["recuperacion"] = path
    return items


def insert_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Cm(15.5))
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_cover(doc):
    for _ in range(5): doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("Instituto de Formación Técnica Superior 24")
    r.font.name = "Times New Roman"; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("Tecnicatura Superior en Desarrollo de Software"); r.font.name = "Times New Roman"; r.font.size = Pt(14)
    for _ in range(5): doc.add_paragraph("")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("Proyecto Final Integrador"); r.font.name = "Times New Roman"; r.font.size = Pt(18); r.font.bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("DVT — Software de Gestión Hotelera"); r.font.name = "Times New Roman"; r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
    for _ in range(7): doc.add_paragraph("")
    for label, value in [("Autor", "Iván Federico Ibarra"), ("Tutor", "Ing. Julio Cesar Leppen"), ("Año", "2026")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(f"{label}: {value}"); r.font.name = "Times New Roman"; r.font.size = Pt(14)
    doc.add_page_break()


def add_front_matter(doc):
    add_title(doc, "Agradecimientos y dedicatorias")
    add_text(doc, [
        "Dedico este proyecto final a mi abuela, Dalmira del Valle Tapia. Su amor, sus enseñanzas y la constancia que transmitió forman parte de cada etapa de este logro.",
        "Agradezco a mi familia por el acompañamiento permanente durante la carrera, al IFTS 24 por brindar el espacio de formación y a cada docente que contribuyó a mi desarrollo profesional y personal. Extiendo un agradecimiento especial al tutor por la orientación brindada en la integración del proyecto.",
    ])
    add_title(doc, "Resumen")
    add_text(doc, [
        "Este Trabajo Final Integrador presenta DVT — Software de Gestión Hotelera, una aplicación web orientada a digitalizar procesos operativos de pequeños y medianos alojamientos. El problema abordado es la dispersión de la información y la dependencia de registros manuales para administrar habitaciones, huéspedes, reservas, consumos y pagos, situaciones que favorecen errores, demoras y dificultades para consultar el historial operativo.",
        "La solución se desarrolló con una arquitectura web de dos capas visibles: un sitio comercial estático y una aplicación interna de operación. El frontend utiliza Vue 3, Vite y Vue Router; el backend utiliza Node.js y Express, organizado en rutas, controladores, servicios y modelos. La persistencia se resuelve con Cloud Firestore mediante Firebase Admin SDK. La seguridad contempla autenticación con JSON Web Tokens, contraseñas protegidas con bcrypt, autorización por roles, variables de entorno y una lista blanca de orígenes CORS.",
        "Como resultado se implementaron módulos de usuarios, habitaciones, huéspedes, reservas, disponibilidad, consumos, pagos y recuperación interna de credenciales. La aplicación se encuentra desplegada como Web Service en Render y dispone de un mecanismo reproducible de datos ficticios para Firestore. Se concluye que la centralización de estos procesos permite mejorar la trazabilidad operativa y ofrece una base técnica mantenible para futuras ampliaciones.",
    ])
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0); p.add_run("Palabras clave: ").bold = True; p.add_run("gestión hotelera, reservas, aplicación web, Firestore, Vue, Express.")
    add_title(doc, "Abstract")
    add_text(doc, [
        "This Final Integrative Project presents DVT — Hotel Management Software, a web application designed to digitize operational processes in small and medium-sized accommodation businesses. The project addresses the fragmentation of information and the reliance on manual records for rooms, guests, reservations, additional charges and payments.",
        "The solution combines a public commercial site with an internal operational application. The frontend uses Vue 3, Vite and Vue Router, while the backend uses Node.js and Express, organized into routes, controllers, services and models. Cloud Firestore is used for persistence through Firebase Admin SDK. Security measures include JSON Web Tokens, bcrypt password hashing, role-based authorization, environment variables and an allowlist for CORS origins.",
        "The resulting system implements users, rooms, guests, reservations, availability, additional charges, payments and an internal temporary-password workflow. It is deployed as a Render Web Service and includes reproducible fictitious Firestore seed data. The project demonstrates that a centralized web solution can improve operational traceability and provide a maintainable foundation for future growth.",
    ])
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0); p.add_run("Keywords: ").bold = True; p.add_run("hotel management, reservations, web application, Firestore, Vue, Express.")
    add_title(doc, "Declaración de ética y consentimiento informado")
    add_text(doc, [
        "Se declara que el presente trabajo fue elaborado con fines académicos y con criterio de honestidad intelectual. La solución presentada fue desarrollada para demostrar competencias de análisis, diseño, construcción, pruebas y despliegue de software. Las explicaciones técnicas se sustentan en el código fuente, en pruebas realizadas sobre el sistema y en fuentes bibliográficas identificadas.",
        "Las capturas, diagramas y datos empleados en el documento no exponen credenciales, claves, tokens ni información personal real de huéspedes. Para los ejemplos reproducibles de base de datos se utilizan nombres, correos y documentos explícitamente ficticios. Esta decisión reduce el riesgo de exposición de información sensible y es coherente con el principio de minimización de datos.",
        "Se autoriza al IFTS 24 a consultar este documento y el material audiovisual asociado exclusivamente con fines educativos y de evaluación. La disponibilidad pública del código y del despliegue queda sujeta a las decisiones del autor y a la configuración de los servicios utilizados.",
    ])
    add_title(doc, "Tabla de contenido")
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0); add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    add_title(doc, "Índice de tablas y figuras")
    p = doc.add_paragraph("Índice de tablas", style="Heading 2"); p.paragraph_format.page_break_before = False
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0); add_field(p, 'TOC \\h \\z \\c "Tabla"')
    p = doc.add_paragraph("Índice de figuras", style="Heading 2"); p.paragraph_format.page_break_before = False
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0); add_field(p, 'TOC \\h \\z \\c "Figura"')
    add_title(doc, "Lista de abreviaturas y siglas")
    add_table(doc, ["Sigla", "Significado"], [
        ("API", "Application Programming Interface; interfaz para la comunicación entre componentes de software."),
        ("CORS", "Cross-Origin Resource Sharing; política que controla orígenes autorizados para solicitudes web."),
        ("EDT", "Estructura de Desglose del Trabajo."),
        ("JWT", "JSON Web Token; formato compacto firmado para transmitir claims de una sesión."),
        ("PMS", "Property Management System; sistema de gestión de propiedades o alojamientos."),
        ("SPA", "Single-Page Application; aplicación que actualiza vistas sin recargar todo el documento."),
        ("UI/UX", "Interfaz de usuario / experiencia de usuario."),
    ], widths=[3, 13], caption="Tabla 1. Abreviaturas empleadas en el documento.")
    add_title(doc, "Glosario de términos clave")
    add_table(doc, ["Término", "Definición aplicada al proyecto"], [
        ("Cuenta corriente", "Acumulación de alojamiento y consumos incluidos en la cuenta de una reserva, menos los pagos confirmados."),
        ("Consumo pendiente", "Cargo operativo que aún puede ser corregido y que no integra el saldo hasta ser incluido en cuenta."),
        ("Hash de contraseña", "Resultado no reversible que permite validar una clave sin almacenarla en texto plano."),
        ("Middleware", "Función ejecutada entre la llegada de una solicitud HTTP y el controlador, utilizada para autenticación, roles y errores."),
        ("Snapshot", "Copia contextual de datos de huésped y habitación dentro de una reserva, pago o consumo para preservar el historial."),
        ("Web Service", "Servicio de Render destinado a aplicaciones dinámicas que responden solicitudes HTTP públicas."),
    ], widths=[4, 12], caption="Tabla 2. Glosario técnico y funcional.")


def add_body(doc, figures):
    add_title(doc, "1. Introducción")
    add_text(doc, [
        "La operación de un alojamiento reúne actividades que ocurren de manera simultánea: asignación de habitaciones, ingreso y salida de huéspedes, registro de servicios adicionales, consulta de disponibilidad y cobro. Cuando estos procesos se resuelven en papeles, planillas aisladas o mensajes informales, se vuelve difícil conocer el estado real del establecimiento y auditar decisiones tomadas durante una estadía.",
        "DVT se concibe como una herramienta web para centralizar esa información en un entorno accesible desde navegador. La propuesta no intenta reemplazar todos los sistemas de una cadena hotelera; se concentra en necesidades frecuentes de alojamientos pequeños y medianos: operar reservas sin superposiciones, conservar datos de huéspedes, registrar consumos con criterio uniforme y conocer el saldo de una estadía.",
        "El documento describe el problema, el alcance, el proceso de desarrollo, la solución técnica, los resultados de las verificaciones y las limitaciones detectadas. Asimismo, establece evidencia suficiente para que la solución pueda ser evaluada y reproducida en un entorno controlado.",
    ])
    add_title(doc, "2. Marco conceptual y antecedentes")
    add_text(doc, [
        "La transformación digital del turismo modificó tanto la relación con el cliente como la organización interna de los prestadores. Buhalis y Law (2008) identificaron a las tecnologías de información como un componente estructural de la gestión turística; más recientemente, Gössling (2021) señaló que la tecnología continúa redefiniendo prácticas, capacidades y expectativas dentro del sector.",
        "En un hotel, un PMS permite articular datos de ocupación, huéspedes y transacciones. El presente proyecto adopta un subconjunto funcional de ese concepto: no incorpora un motor de reservas público ni integración contable externa, pero resuelve el circuito interno de recepción. Esta delimitación permite concentrar el esfuerzo académico en reglas de negocio verificables y en una arquitectura web mantenible.",
        "Desde el punto de vista técnico, una interfaz basada en componentes facilita representar vistas independientes y reactivas. Vue se define como un framework de interfaz declarativo y basado en componentes (Vue.js, 2026). Vite complementa este enfoque con un servidor de desarrollo y un proceso de empaquetado para producción (Vite, 2026).",
    ])
    add_title(doc, "3. Planteamiento del problema")
    add_text(doc, [
        "El problema principal consiste en la falta de una fuente única de información para la recepción de un hotel. Una reserva puede ser anotada en un medio, el consumo registrado en otro y el pago informado sin relación explícita con el huésped o la habitación. Esta fragmentación dificulta responder preguntas básicas: qué habitación está disponible, qué servicios se cargaron a una estadía y cuál es el saldo pendiente.",
        "La ausencia de validaciones sistemáticas también incrementa el riesgo de sobreventa. Si dos personas registran reservas sobre la misma habitación y fechas sin una comprobación común, el conflicto se detecta tarde. Asimismo, una descripción libre de gastos provoca inconsistencia al momento de responder un reclamo del huésped o preparar un cierre operativo.",
        "DVT aborda estas situaciones con reglas de negocio en el backend, no sólo con controles visuales. Las reservas activas se comparan por rango de fechas; los consumos se clasifican en categorías normalizadas; los pagos se vinculan a una reserva existente; y las acciones sensibles se restringen según rol.",
    ])
    add_title(doc, "4. Justificación")
    add_text(doc, [
        "La justificación del proyecto se vincula con la necesidad de ofrecer una solución de bajo costo inicial, accesible desde navegador y sin instalación de infraestructura local. Una aplicación web hospedada permite centralizar el acceso y reduce la dependencia de archivos locales dispersos. El uso de servicios administrados también disminuye la carga operativa asociada a mantener un servidor propio durante el alcance académico.",
        "La utilidad no se limita a almacenar registros. Al utilizar snapshots de huésped y habitación dentro de reservas, consumos y pagos, se preserva el contexto histórico aunque luego se modifique el registro maestro. Este aspecto es relevante frente a consultas posteriores y contribuye a la trazabilidad del circuito de recepción.",
        "Finalmente, el proyecto integra contenidos de la tecnicatura: programación web, APIs REST, bases de datos NoSQL, seguridad, experiencia de usuario, despliegue, documentación y aseguramiento de calidad.",
    ])
    add_title(doc, "5. Objetivos")
    p = doc.add_paragraph("5.1 Objetivo general", style="Heading 2"); p.paragraph_format.page_break_before = False
    add_text(doc, ["Desarrollar y desplegar una aplicación web de gestión hotelera que centralice procesos operativos de recepción, con controles de acceso, persistencia en la nube y una interfaz responsiva para usuarios administrativos."])
    p = doc.add_paragraph("5.2 Objetivos específicos", style="Heading 2"); p.paragraph_format.page_break_before = False
    add_bullets(doc, [
        "Implementar una API REST modular con reglas de negocio para habitaciones, huéspedes, reservas, consumos, pagos y usuarios.",
        "Evitar solapamientos de reservas mediante validación del rango de fechas y del estado operativo.",
        "Construir una interfaz SPA que permita consultar, crear y actualizar información según el rol autenticado.",
        "Aplicar autenticación JWT, hash de contraseñas, autorización por rol, CORS y gestión de secretos por variables de entorno.",
        "Desplegar el sistema en Render y documentar una forma segura de reproducir datos ficticios en Firestore.",
    ])
    add_title(doc, "6. Alcance y restricciones")
    p = doc.add_paragraph("6.1 Alcance funcional", style="Heading 2"); p.paragraph_format.page_break_before = False
    add_text(doc, [
        "El alcance incluye un sitio comercial público y una aplicación interna. La aplicación administra usuarios internos, habitaciones, huéspedes, reservas, operaciones de check-in y check-out, consumos adicionales, pagos y visualización de datos operativos. El sistema contempla dos roles: administrador y recepcionista.",
        "La recuperación de acceso es un flujo interno: el usuario solicita una clave temporal desde el login y un administrador revisa la solicitud antes de asignarla. La clave temporal exige cambio al próximo inicio de sesión. El diseño evita prometer un envío de contraseña por correo que no esté efectivamente configurado.",
    ])
    p = doc.add_paragraph("6.2 Restricciones y exclusiones", style="Heading 2"); p.paragraph_format.page_break_before = False
    add_bullets(doc, [
        "No se implementa un motor de reserva público ni pasarela de pago electrónica; los pagos son registros operativos confirmados por recepción.",
        "No se integra facturación fiscal, contabilidad externa, cerraduras inteligentes ni channel manager.",
        "No se utiliza información real de huéspedes en las demostraciones ni en el seed del repositorio.",
        "El plan gratuito de Render puede suspender una instancia inactiva; el servicio se reactiva con una nueva solicitud, por lo que puede ocurrir una demora inicial.",
        "La aplicación está dirigida a navegadores modernos; no se declara soporte específico para navegadores obsoletos.",
    ])
    add_title(doc, "7. Criterios de aceptación")
    add_table(doc, ["ID", "Criterio verificable", "Evidencia esperada"], [
        ("CA-01", "Una reserva activa no puede superponerse con otra de la misma habitación.", "Respuesta de conflicto y persistencia sin duplicidad de fechas."),
        ("CA-02", "Los usuarios no autenticados no acceden a módulos privados.", "Respuesta HTTP 401 sin token válido."),
        ("CA-03", "El recepcionista puede operar reservas, consumos y pagos sin administrar usuarios.", "Matrices de ruta y pruebas por rol."),
        ("CA-04", "Un consumo pendiente sólo integra el saldo al ser incluido en cuenta.", "Cambio de estado a EN_CUENTA y cálculo de saldo."),
        ("CA-05", "Un pago confirmado se relaciona con una reserva existente.", "Registro en PAGOS con snapshot de reserva."),
        ("CA-06", "El sistema se entrega accesible desde una URL pública y con health check.", "URL de Render y GET /api/health."),
    ], widths=[2, 8, 6], caption="Tabla 3. Criterios de aceptación funcionales y técnicos.")
    add_title(doc, "8. Metodología de desarrollo")
    add_text(doc, [
        "Se aplicó un enfoque incremental. Primero se revisó una base de backend existente y se identificaron módulos, rutas y modelos. Luego se construyó la interfaz interna, se integraron endpoints, se validaron reglas de negocio y se ajustó la experiencia de uso a partir de pruebas manuales con los roles administrativos.",
        "Cada incremento se trató como una unidad verificable: autenticación, gestión de habitaciones, reservas, consumos, pagos, filtros y recuperación de credenciales. La revisión continua permitió corregir problemas de relación entre estados. Por ejemplo, se separó la inclusión de un consumo en la cuenta del acto de confirmar un pago, evitando que el botón de facturar se interprete como cobro de dinero.",
        "La fase final incorpora revisión de código, build del frontend, pruebas de rutas protegidas, validación del despliegue y documentación. El enfoque es adecuado para un proyecto académico porque facilita vincular cada requerimiento con una evidencia observable.",
    ])
    add_title(doc, "9. Estructura de Desglose del Trabajo")
    add_table(doc, ["Nivel", "Componente", "Entregable"], [
        ("1", "Planificación", "Problema, objetivos, alcance, riesgos y cronograma."),
        ("2", "Diseño", "Arquitectura, modelo de datos, flujos y diseño responsivo."),
        ("3", "Backend", "API Express, Firestore, autenticación, roles y reglas de negocio."),
        ("4", "Frontend", "Landing pública, SPA Vue, rutas y vistas operativas."),
        ("5", "Calidad", "Pruebas manuales, lint, build y corrección de incidencias."),
        ("6", "Despliegue y entrega", "Render, repositorio, seed, tesis, presentación y video."),
    ], widths=[2, 5, 9], caption="Tabla 4. EDT resumida del proyecto.")
    add_text(doc, ["La EDT relaciona el trabajo técnico con los entregables exigidos para el pre-final. Su utilidad consiste en evitar que el despliegue, la documentación y la evidencia de pruebas queden como actividades separadas del desarrollo funcional."])
    add_title(doc, "10. Cronograma de alto nivel")
    add_table(doc, ["Etapa", "Resultado", "Estado al pre-final"], [
        ("Análisis y alcance", "Objetivos, criterios y restricciones definidos.", "Completado."),
        ("Backend y Firestore", "Módulos y reglas de negocio implementados.", "Completado con revisión final."),
        ("Frontend y responsividad", "Landing y aplicación interna integradas.", "Completado con validación visual pendiente."),
        ("Despliegue", "Web Service Render y health check configurados.", "Completado; requiere chequeo previo a defensa."),
        ("Documentación", "Tesis, anexos, README y seed de datos ficticios.", "En consolidación."),
        ("Presentación y video", "PPTX y grabación de 15 a 30 minutos.", "Pendiente de producción."),
    ], widths=[4, 7, 5], caption="Tabla 5. Cronograma y seguimiento de entregables.")
    add_text(doc, ["El cronograma debe actualizarse con fechas reales antes de la entrega final. La prioridad de la última semana es consolidar evidencias reproducibles y ensayar el recorrido de demostración con un plan alternativo ante una demora de arranque del hosting gratuito."])
    add_title(doc, "11. Arquitectura de la solución")
    add_text(doc, [
        "La solución adopta una arquitectura cliente-servidor. La landing comercial se distribuye como contenido estático; la aplicación interna se compila con Vite y se publica bajo la ruta /app. Ambas son servidas por el mismo proceso Express en producción, por lo que la navegación de la aplicación evita depender de un segundo hosting.",
        "La API recibe solicitudes HTTP, aplica middlewares globales y monta rutas por módulo. El backend concentra la autenticación, la autorización y las validaciones de negocio antes de acceder a Firestore. Esta separación evita delegar decisiones críticas exclusivamente en el navegador y permite reutilizar las reglas desde diferentes vistas.",
        "Render ejecuta el build definido en render.yaml, inicia Node.js y consulta /api/health para verificar disponibilidad. La configuración sensible se carga desde variables de entorno del panel de Render y no desde el repositorio.",
    ])
    insert_figure(doc, figures["arquitectura"], "Figura 1. Arquitectura de solución técnica de DVT.")
    add_title(doc, "12. Backend y organización por capas")
    add_text(doc, [
        "El backend utiliza Node.js, un entorno de ejecución de JavaScript diseñado para operaciones de entrada y salida no bloqueantes (Node.js, 2026). Express organiza el servicio HTTP y permite aplicar middlewares antes de los controladores. El archivo index.js se limita a configurar el servidor, CORS, parseo JSON, rutas, archivos estáticos y manejo de errores.",
        "Cada módulo sigue una organización de responsabilidades: las rutas declaran endpoint y permisos; los controladores traducen la solicitud HTTP; los servicios aplican validaciones y reglas; y los modelos acceden a las colecciones de Firestore. Esta estructura reduce acoplamiento y facilita explicar dónde se valida cada regla durante una defensa.",
        "Los módulos principales son habitaciones, huéspedes, reservas, disponibilidades, consumos, pagos, usuarios y recuperación de credenciales. Además existe un servicio operacional de reservas que compone alojamiento, consumos cobrables, pagos confirmados y saldo pendiente sin modificar el historial original.",
    ])
    add_title(doc, "13. Frontend, navegación y responsividad")
    add_text(doc, [
        "El frontend interno se desarrolló como SPA con Vue 3 y Vue Router. Vue utiliza un modelo de componentes y reactividad que permite actualizar la interfaz en función de los datos obtenidos desde la API (Vue.js, 2026). Cada módulo operativo se representa en una vista y utiliza un servicio común para adjuntar el token JWT a las solicitudes.",
        "Vite proporciona el servidor de desarrollo y genera el bundle productivo de la aplicación (Vite, 2026). La configuración establece base /app/, consistente con el punto de montaje que Express utiliza en producción. Esta decisión permite conservar la landing comercial en la raíz y dirigir al usuario autenticado hacia la aplicación interna.",
        "La responsividad se trabajó mediante grillas flexibles, adaptación del sidebar, formularios de ancho disponible y representación en tarjetas de información que no resulta legible en una tabla estrecha. La calidad de la interfaz se evalúa en resoluciones de escritorio, tableta y teléfono antes de la demostración final.",
    ])
    add_title(doc, "14. Persistencia y modelo de datos")
    add_text(doc, [
        "Cloud Firestore es una base NoSQL orientada a documentos organizados en colecciones (Firebase, 2026a). El proyecto usa las colecciones USUARIOS, HABITACIONES, HUESPEDES, RESERVAS, DISPONIBILIDADES, CONSUMOS_EXTRAS, PAGOS y SOLICITUDES_RECUPERACION_CLAVE. La elección permite representar datos operativos y objetos anidados sin requerir un esquema relacional rígido.",
        "Las reservas guardan snapshots del huésped y de la habitación. Pagos y consumos conservan además un snapshot de la reserva. De este modo, si se modifica un nombre, una tarifa o una habitación maestra, los movimientos históricos siguen mostrando el contexto que tenían al registrarse. Esta es una decisión de trazabilidad, no una duplicación accidental.",
        "DISPONIBILIDADES almacena la proyección diaria de una habitación con un identificador determinístico compuesto por habitación y fecha. La reserva es la fuente operativa que bloquea o libera esa proyección. El servicio también vuelve a consultar reservas activas para prevenir solapamientos.",
    ])
    insert_figure(doc, figures["datos"], "Figura 2. Modelo lógico simplificado de las colecciones principales.")
    add_title(doc, "15. Módulo de autenticación y sesiones")
    add_text(doc, [
        "El login recibe correo y contraseña, consulta el usuario por email y compara la clave con el hash almacenado. Cuando las credenciales son válidas, el backend firma un JWT con identificador, correo y rol, con vencimiento de ocho horas. El token se envía en el encabezado Authorization bajo el esquema Bearer para acceder a rutas privadas.",
        "JWT es un estándar para representar claims en un formato compacto y firmado (Jones, Bradley, & Sakimura, 2015). En DVT el token no contiene la contraseña ni datos innecesarios; sólo incluye la información requerida por el middleware para reconocer una sesión y aplicar roles.",
        "Las rutas privadas rechazan solicitudes sin token o con token vencido mediante respuesta HTTP 401. Luego, el middleware de roles controla que el rol del token pertenezca a los permitidos por la operación. Esta segunda verificación es necesaria porque estar autenticado no implica poder administrar funciones críticas.",
    ])
    add_title(doc, "16. Recuperación interna de credenciales")
    add_text(doc, [
        "El sistema no entrega contraseñas por correo ni expone si un correo existe en la base. Desde la pantalla de acceso, una persona puede solicitar una clave temporal. El backend devuelve una respuesta genérica y, sólo si el email corresponde a un usuario, registra una solicitud PENDIENTE asociada a esa cuenta.",
        "Un administrador visualiza las solicitudes pendientes en el módulo Usuarios, verifica la identidad por un canal organizacional y asigna una clave temporal. Se persiste exclusivamente su hash bcrypt, el usuario queda marcado con mustChangePassword y la solicitud registra quién la resolvió. La clave temporal no se guarda en texto plano en el registro de solicitud.",
        "Este flujo es consistente con el alcance institucional del sistema: la recuperación queda bajo control del administrador de la organización y evita simular un proveedor de correo que no forme parte del despliegue. También crea una evidencia de auditoría para responder ante una consulta sobre el cambio de acceso.",
    ])
    insert_figure(doc, figures["recuperacion"], "Figura 3. Flujo de recuperación interna de credenciales.")
    add_title(doc, "17. Módulo de usuarios y permisos")
    add_text(doc, [
        "El administrador puede registrar usuarios internos, modificar nombre, rol y estado, y asignar contraseñas temporales. Los roles válidos se normalizan en el servicio de usuarios y son exclusivamente ADMIN y RECEPCIONISTA. Esta decisión evita roles declarados en la documentación que no estén respaldados por las reglas del backend.",
        "La matriz de permisos aplica el principio de mínimo privilegio: el recepcionista puede ejecutar tareas de recepción, pero no crear usuarios, eliminar pagos, borrar consumos ni gestionar habitaciones. El administrador conserva esas funciones de mayor impacto.",
    ])
    add_table(doc, ["Funcionalidad", "ADMIN", "RECEPCIONISTA"], [
        ("Usuarios y claves temporales", "Crear, editar, desactivar y asignar clave.", "Sin acceso."),
        ("Habitaciones", "Alta, edición, baja y consulta.", "Consulta."),
        ("Huéspedes y reservas", "Operación completa.", "Operación diaria."),
        ("Consumos", "Crear, editar pendientes, incluir, cerrar o eliminar.", "Crear, editar pendientes e incluir en cuenta."),
        ("Pagos", "Crear, consultar, editar o eliminar.", "Crear y consultar."),
    ], widths=[5, 5.5, 5.5], caption="Tabla 6. Matriz de permisos implementada.")
    add_title(doc, "18. Módulo de habitaciones y huéspedes")
    add_text(doc, [
        "Las habitaciones se registran con número, tipo, tarifa por noche, capacidad, estado, piso, descripción, amenities y fotografías. El backend valida campos obligatorios, tarifa positiva, capacidad entera y estados permitidos: DISPONIBLE, OCUPADA o MANTENIMIENTO. Una habitación en mantenimiento no puede recibir una reserva activa.",
        "El módulo de huéspedes conserva nombre, apellido, documento, teléfono, correo, nacionalidad y observaciones. La validación se realiza antes de crear una reserva, por lo que la reserva puede vincularse a un huésped existente o crear uno nuevo con datos completos. Esto reduce la repetición de carga y mantiene una entidad consultable.",
        "La separación entre habitación y huésped permite que una misma persona tenga más de una reserva histórica y que una habitación sea utilizada por distintos huéspedes a lo largo del tiempo. Las referencias se complementan con snapshots para conservar la lectura histórica.",
    ])
    add_title(doc, "19. Módulo de reservas y disponibilidad")
    add_text(doc, [
        "Al crear una reserva se validan habitación, fecha de inicio y fecha de fin. La fecha de entrada debe ser anterior a la salida. El servicio consulta las reservas de la habitación y compara intervalos de las reservas con estado pending, confirmed o checked_in. Las reservas canceladas o finalizadas no bloquean futuras disponibilidades.",
        "Una reserva recibe un código legible para recepción y conserva snapshots de huésped y habitación. Durante la operación se admiten estados pending, confirmed, checked_in, checked_out y cancelled. El check-in sólo es válido desde la fecha de entrada; el check-out sólo se permite luego de un check-in. La cancelación es una baja lógica que conserva historial y libera fechas.",
        "La vista de reservas ofrece filtros por estado operativo, mes de ingreso, huésped y habitación. Al limpiar filtros se vuelve a mostrar el conjunto correspondiente a la vista activa, conducta relevante para la eficiencia de la recepción al consultar historial o reservas pendientes.",
    ])
    add_title(doc, "20. Módulo de consumos y pagos")
    add_text(doc, [
        "Los consumos extras representan servicios o cargos vinculados a una reserva activa. El sistema normaliza las categorías gastronomía, lavandería, estacionamiento, spa y otros, evitando categorías ambiguas o eliminadas del alcance. La descripción sigue siendo necesaria para especificar el detalle, por ejemplo, desayuno continental o lavado y planchado de prendas.",
        "Un consumo nuevo se registra como PENDIENTE_FACTURACION. En ese estado puede ser corregido por el personal autorizado. Al confirmar que debe formar parte de la cuenta, la acción lo cambia a EN_CUENTA; recién entonces se incorpora al saldo a cobrar. Los estados FACTURADO y CERRADO se conservan como compatibilidad histórica y no se editan.",
        "El pago es un movimiento distinto. Para crear un pago, el backend exige una reserva existente, monto positivo y método permitido: efectivo, tarjeta o transferencia. El pago se registra como CONFIRMADO y reduce el saldo de la reserva. Esta separación evita afirmar que un consumo fue cobrado sólo por haber sido agregado a la cuenta.",
    ])
    insert_figure(doc, figures["consumos"], "Figura 4. Secuencia de consumo, cuenta corriente y pago.")
    add_title(doc, "21. Panel operativo y filtros")
    add_text(doc, [
        "El panel operativo concentra métricas de reservas activas, entradas y salidas del día, habitaciones ocupadas y próximas reservas. Las métricas se calculan desde la información entregada por la API y permiten a recepción priorizar acciones de la jornada.",
        "Los módulos de reservas, pagos y consumos incluyen filtros con criterios funcionales: estado, período, huésped, habitación o reserva. El principio de diseño aplicado es mostrar primero un conjunto útil de registros y permitir acotar gradualmente. Todos los filtros disponen de una acción para limpiar la selección y restablecer la vista.",
        "Para pantallas angostas, las listas extensas se presentan mediante tarjetas o contenedores adaptables en lugar de obligar a leer una tabla comprimida. Esta decisión preserva la lectura de identificadores, importes, estados y acciones en dispositivos de uso cotidiano.",
    ])
    add_title(doc, "22. Evaluación de riesgos y seguridad")
    add_text(doc, [
        "La seguridad se consideró como un conjunto de controles técnicos y operativos. OWASP (2026) recomienda evitar el almacenamiento de contraseñas en texto plano y utilizar algoritmos lentos y adaptativos. DVT utiliza bcrypt con factor de trabajo 10 mediante bcryptjs; cada hash se guarda en Firestore y nunca se devuelve al frontend.",
        "Los secretos se leen desde variables de entorno. En producción, Firebase se configura mediante una cuenta de servicio codificada en Base64 o JSON gestionada en Render; el archivo local firebase-credentials.json está excluido del repositorio. CORS utiliza una lista blanca de orígenes, y JWT_SECRET es obligatorio para firmar y validar sesiones.",
        "La seguridad no se declara como absoluta. La falta de autenticación multifactor, límites explícitos de intentos de login y pruebas automatizadas de seguridad constituye una oportunidad de mejora. Estas limitaciones se documentan para no sobreestimar el alcance de una solución académica.",
    ])
    add_table(doc, ["Riesgo", "Impacto", "Control actual", "Mejora futura"], [
        ("Exposición de credenciales", "Acceso no autorizado a Firestore.", "Variables de entorno y .gitignore.", "Rotación y auditoría de secretos."),
        ("Acceso indebido por rol", "Cambio o borrado de información.", "JWT y middleware de roles.", "Pruebas automatizadas por permiso."),
        ("Contraseña comprometida", "Suplantación de usuario.", "bcrypt y clave temporal con cambio obligatorio.", "MFA y límite de intentos."),
        ("Caída o demora de hosting", "Demora de acceso durante demostración.", "Health check y despliegue versionado.", "Plan pago o monitoreo externo."),
        ("Error de operación", "Cargo o reserva incorrectos.", "Validaciones, estados y snapshots.", "Auditoría ampliada y confirmaciones adicionales."),
    ], widths=[3.3, 3.3, 4.7, 4.7], caption="Tabla 7. Riesgos, controles y oportunidades de mejora.")
    add_title(doc, "23. Despliegue en Render")
    add_text(doc, [
        "El despliegue se realiza en Render como Web Service porque el proyecto ejecuta código del lado servidor con Express. Render describe este tipo de servicio como adecuado para aplicaciones dinámicas accesibles por HTTP y permite desplegar desde una rama conectada de Git (Render, 2026a). El repositorio incorpora render.yaml para dejar declarados buildCommand, startCommand, versión de Node y healthCheckPath.",
        "Durante el build se instalan dependencias del backend, luego las del frontend y se genera DVT-SoftwareHotelero/dist. Al iniciar, Express sirve tanto los activos de la landing como el bundle de Vue. El endpoint GET /api/health responde un objeto JSON con estado, nombre de servicio y timestamp; constituye una evidencia simple de que el proceso y la API están disponibles.",
        "El hosting gratuito presenta una limitación conocida: Render puede detener instancias sin actividad y reanudarlas al recibir una nueva solicitud (Render, 2026b). Para la exposición se recomienda abrir la URL con anticipación, confirmar el health check y mantener una grabación breve de respaldo del recorrido funcional.",
    ])
    add_title(doc, "24. Reproducción de base de datos")
    add_text(doc, [
        "Una condición de reproducibilidad es poder demostrar el sistema sin depender de la base utilizada durante el desarrollo. Por este motivo el repositorio incorpora database/seed-data.example.json y scripts/seed-firestore.js. El JSON contiene exclusivamente registros ficticios de habitaciones, huésped, reserva, consumos y pagos; no incluye credenciales, información personal real ni exportaciones de producción.",
        "El script requiere el argumento explícito --apply y dos variables de entorno para las claves de usuarios de demostración. Antes de crear cada documento verifica que el identificador use el prefijo demo- y omite documentos existentes. No elimina registros ni sobrescribe información operativa. De este modo, la inicialización es idempotente para una base nueva y prudente frente a un entorno con datos.",
        "La documentación de la carpeta database explica el procedimiento y las colecciones resultantes. La reproducibilidad debe ejecutarse sólo en un proyecto Firebase de demostración o desarrollo, nunca sobre una base utilizada por un establecimiento real sin un plan de respaldo y autorización.",
    ])
    add_title(doc, "25. Estrategia de pruebas")
    add_text(doc, [
        "La estrategia de pruebas combina verificaciones estáticas, pruebas manuales de integración y comprobaciones de disponibilidad. En el frontend se ejecuta ESLint y el build de Vite. En el backend se verifica sintaxis de archivos críticos y se realizan solicitudes a endpoints protegidos para confirmar que, sin token, responden HTTP 401.",
        "Las pruebas manuales se organizan por flujo y no sólo por pantalla: ingreso como administrador y recepcionista, alta de habitación, reserva con huésped, intento de solapamiento, check-in, registro de consumo, incorporación a cuenta, registro de pago, filtros, solicitud de clave temporal y cambio obligado de contraseña.",
        "Para el pre-final se propone registrar en una planilla de evidencia la fecha, caso, usuario de prueba, resultado esperado, resultado obtenido y captura asociada. Esta práctica permite separar lo efectivamente validado de las mejoras que queden planificadas.",
    ])
    add_table(doc, ["Caso", "Acción", "Resultado esperado"], [
        ("PF-01", "Acceder a /api/reservas sin JWT.", "HTTP 401; no se filtran datos."),
        ("PF-02", "Crear dos reservas activas solapadas para una habitación.", "La segunda operación se rechaza con conflicto."),
        ("PF-03", "Registrar consumo pendiente e incluirlo en cuenta.", "Cambia de PENDIENTE_FACTURACION a EN_CUENTA y afecta saldo."),
        ("PF-04", "Registrar un pago sobre una reserva válida.", "Se crea movimiento CONFIRMADO y reduce saldo."),
        ("PF-05", "Solicitar clave temporal y resolver como ADMIN.", "Se registra auditoría y se exige cambio de contraseña."),
        ("PF-06", "Abrir URL de Render y /api/health.", "Landing, login y estado responden correctamente."),
    ], widths=[2, 7, 7], caption="Tabla 8. Casos de prueba prioritarios para el pre-final.")
    add_title(doc, "26. Resultados de la validación")
    add_text(doc, [
        "La revisión técnica confirmó que el frontend compila y que el análisis ESLint finaliza sin errores. Los archivos de servidor y el script de seed superan validación sintáctica. La ejecución del seed sin --apply se comporta como modo informativo, por lo que no solicita credenciales ni realiza escrituras accidentales.",
        "La verificación de rutas privadas sin token confirmó respuestas HTTP 401 para módulos de huéspedes, habitaciones, reservas, usuarios, pagos y consumos. El endpoint público /api/health, la landing y la ruta /app/login respondieron satisfactoriamente en la URL desplegada durante la auditoría del proyecto.",
        "Las pruebas manuales locales permitieron comprobar el flujo de clave temporal, la distinción entre agregar un consumo a la cuenta y registrar un pago, los filtros restablecibles y la visibilidad operativa de consumos para recepción. Estas evidencias deben volver a ejecutarse y registrarse el día previo al pre-final debido a que la información de Firestore y el estado de Render son dinámicos.",
    ])
    add_title(doc, "27. Calidad, usabilidad y responsividad")
    add_text(doc, [
        "La calidad se interpreta en relación con los objetivos del sistema, no sólo con la ausencia de errores. Para recepción, una operación útil debe comunicar claramente qué datos se están registrando, qué estado financiero tiene un consumo y qué acciones son permitidas. Por ese motivo las etiquetas se ajustaron a términos operativos como Agregar a cuenta, Confirmar pago y Limpiar filtros.",
        "El diseño incorpora una jerarquía visual entre el sitio comercial y el sistema interno. La landing comunica propuesta de valor y redirige al acceso; la aplicación prioriza formularios, estados, tablas y acciones operativas. El botón para volver arriba en la landing y la adaptación de navegación en resoluciones reducidas responden a problemas observables de recorrido y lectura.",
        "La evaluación de responsividad no reemplaza un estudio formal de accesibilidad. Como mejora futura se recomienda incorporar pruebas con lectores de pantalla, contraste automatizado, navegación íntegra por teclado y auditorías de rendimiento en dispositivos de menor capacidad.",
    ])
    add_title(doc, "28. Costos directos e indirectos")
    add_table(doc, ["Concepto", "Tipo", "Situación durante el proyecto", "Observación"], [
        ("Desarrollo y documentación", "Indirecto", "Tiempo del autor.", "No se monetiza en la versión académica."),
        ("Node.js, Vue, Vite, Express", "Directo", "Sin licencia comercial obligatoria.", "Tecnologías de código abierto."),
        ("Firebase / Firestore", "Directo variable", "Uso dentro de cuota o configuración de prueba.", "El costo depende de lecturas, escrituras, almacenamiento y tráfico."),
        ("Render", "Directo", "Plan gratuito para demostración.", "Puede requerir plan pago para disponibilidad sostenida."),
        ("Dominio y soporte", "Futuro", "No incluido.", "Requeridos para una operación comercial formal."),
        ("Capacitación del personal", "Indirecto", "No presupuestada.", "Necesaria antes de adopción en un hotel real."),
    ], widths=[3, 2.5, 5.2, 5.3], caption="Tabla 9. Costos directos e indirectos del proyecto.")
    add_text(doc, ["El uso de planes gratuitos es apropiado para la presentación académica, pero no debe confundirse con una garantía de costo cero a escala comercial. Firestore factura según operaciones y almacenamiento, mientras que el tipo de instancia de Render determina disponibilidad y recursos. Por ello, un despliegue productivo debe estimar volumen, soporte, dominio, respaldo y capacitación."])
    add_title(doc, "29. Discusión")
    add_text(doc, [
        "La principal contribución del proyecto es convertir necesidades operativas habituales en reglas explícitas de software. La prevención de solapamientos, la separación entre consumo y pago, la conservación de snapshots y la autorización por rol muestran que la gestión hotelera no se reduce a formularios CRUD. Cada regla representa una decisión que evita inconsistencias en recepción.",
        "La arquitectura elegida favorece la rapidez de implementación y el despliegue. Sin embargo, también presenta decisiones a evaluar para una evolución comercial: Firestore es flexible, pero requiere monitorear el costo de consultas y diseñar índices; Render gratuito facilita la demo, pero el spin-down no es ideal para atención permanente; y bcrypt cumple el almacenamiento de claves actual, aunque una evolución puede considerar Argon2id y autenticación multifactor según el contexto.",
        "El alcance interno de recuperación de acceso es deliberado. En lugar de simular correo electrónico, el sistema registra una solicitud y deja la decisión al administrador. Esta elección es más coherente con la evidencia disponible y permite defender la funcionalidad sin afirmar integraciones inexistentes.",
    ])
    add_title(doc, "30. Conclusiones")
    add_text(doc, [
        "DVT cumple el objetivo de integrar una aplicación web para la gestión operativa básica de un hotel. La solución permite registrar y consultar entidades relacionadas, aplicar reglas de reserva, ordenar consumos y pagos, diferenciar permisos y mantener información en Firestore. El despliegue en Render demuestra que el sistema puede ejecutarse fuera del entorno local y ser accedido mediante una URL pública.",
        "El proyecto también evidencia la importancia de documentar decisiones técnicas. La reproducción de una base ficticia, la descripción de variables de entorno, la matriz de permisos y los casos de prueba convierten la aplicación en un entregable más verificable. La tesis no presenta al sistema como producto comercial terminado, sino como una solución académica funcional con limitaciones explicitadas.",
        "Como líneas futuras se proponen: recuperación mediante canal de correo real con validación adicional, facturación fiscal, reportes financieros, registro de auditoría ampliado, pruebas automatizadas, accesibilidad, integración con pasarela de pago y monitoreo de disponibilidad. Estas ampliaciones deben partir de los controles ya implementados y de un análisis de requerimientos de cada establecimiento.",
    ])
    add_title(doc, "31. Estrategia de difusión y presentación")
    add_text(doc, [
        "La presentación oral debe iniciar con el problema operativo y no con el código. Luego conviene mostrar la landing, iniciar sesión con un usuario de prueba, crear o consultar una reserva, registrar un consumo, incorporarlo a cuenta y mostrar el pago. El cierre debe explicar roles, seguridad, Render y la capacidad de replicar datos ficticios.",
        "El video de entre 15 y 30 minutos puede seguir el mismo hilo: contexto y objetivo; arquitectura; recorrido funcional; seguridad y recuperación de acceso; despliegue; pruebas; límites y trabajo futuro. La cámara encendida permite cumplir la condición institucional, mientras que la captura de pantalla documenta cada acción realizada.",
        "Para reducir contingencias se recomienda mantener abierta la URL de Render antes de comenzar, revisar /api/health, tener una cuenta de demostración disponible y preparar una captura o grabación corta de respaldo en caso de demora transitoria de la instancia gratuita.",
    ])
    add_title(doc, "32. Control de cambios")
    add_table(doc, ["Versión", "Fecha", "Cambio principal", "Responsable"], [
        ("0.1", "2026", "Base inicial de API y entidades operativas.", "Autor."),
        ("0.2", "2026", "Integración de Vue, landing comercial y rutas internas.", "Autor."),
        ("0.3", "2026", "Filtros, responsividad y flujo de consumos/cuenta/pagos.", "Autor."),
        ("0.4", "2026", "Recuperación interna de credenciales, seed ficticio y documentación de despliegue.", "Autor."),
        ("1.0", "Pre-final", "Consolidación de tesis, evidencias, presentación y video.", "Autor y tutoría."),
    ], widths=[2, 2.5, 8.5, 3], caption="Tabla 10. Registro de cambios de alto nivel.")
    add_title(doc, "33. Entregables")
    add_table(doc, ["Entregable", "Ubicación o evidencia", "Estado"], [
        ("Código fuente", "Repositorio GitHub del proyecto.", "Requiere último push validado antes del pre-final."),
        ("Aplicación desplegada", "https://gestionhotelera-dm5p.onrender.com", "Disponible; verificar antes de exposición."),
        ("Health check", "/api/health", "Disponible en el Web Service."),
        ("Datos de demostración", "database/seed-data.example.json y scripts/seed-firestore.js", "Incluido en repositorio local."),
        ("Tesis DOCX", "Este documento y anexos.", "En consolidación."),
        ("Presentación y video", "PPTX y grabación con cámara.", "Pendiente de producción."),
    ], widths=[4, 8, 4], caption="Tabla 11. Entregables exigidos y evidencia asociada.")


def add_references_and_annexes(doc):
    add_title(doc, "Bibliografía")
    refs = [
        "Buhalis, D., & Law, R. (2008). Progress in information technology and tourism management: 20 years on and 10 years after the Internet—The state of eTourism research. Tourism Management, 29(4), 609–623. https://doi.org/10.1016/j.tourman.2008.01.005",
        "Firebase. (2026a). Cloud Firestore. https://firebase.google.com/docs/firestore",
        "Firebase. (2026b). Cloud Firestore data model. https://firebase.google.com/docs/firestore/data-model",
        "Gössling, S. (2021). Tourism, technology and ICT: A critical review of affordances and concessions. Journal of Sustainable Tourism, 29(5), 733–750. https://doi.org/10.1080/09669582.2021.1873353",
        "Jones, M., Bradley, J., & Sakimura, N. (2015). JSON Web Token (JWT) (RFC 7519). Internet Engineering Task Force. https://doi.org/10.17487/RFC7519",
        "Node.js. (2026). Introduction to Node.js. https://nodejs.org/learn",
        "OWASP Foundation. (2026). Password storage cheat sheet. https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html",
        "Render. (2026a). Web services. https://render.com/docs/web-services",
        "Render. (2026b). Your first Render deploy. https://render.com/docs/your-first-deploy",
        "Vite. (2026). Getting started. https://vite.dev/guide/",
        "Vue.js. (2026). Introduction. https://vuejs.org/guide/introduction",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.line_spacing = 2
    add_title(doc, "Anexo A. Arquitectura y estructura de módulos")
    add_text(doc, [
        "La figura siguiente resume la distribución de responsabilidades que se utiliza para explicar el sistema durante la defensa. El repositorio incluye la separación en carpetas de módulos, el frontend Vue, la landing pública y los archivos de despliegue.",
    ])
    diagram = """Cliente web\n  ├── Landing pública (public/)\n  └── SPA interna (DVT-SoftwareHotelero/)\n        └── API REST (index.js + src/)\n              ├── Rutas y middlewares\n              ├── Controladores\n              ├── Servicios y reglas de negocio\n              └── Modelos Firestore\n"""
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.left_indent = Cm(1.27)
    r = p.add_run(diagram); r.font.name = "Courier New"; r.font.size = Pt(10)
    add_title(doc, "Anexo B. Colecciones y campos principales")
    add_table(doc, ["Colección", "Propósito", "Campos representativos"], [
        ("USUARIOS", "Cuentas internas y roles.", "nombre, email, password hash, rol, activo, mustChangePassword."),
        ("HABITACIONES", "Inventario hotelero.", "numero, tipo, precio, capacidad, estado, amenities."),
        ("HUESPEDES", "Datos de huéspedes.", "nombre, apellido, documento, teléfono, email."),
        ("RESERVAS", "Estadías y contexto histórico.", "codigo, fechas, estado, snapshots, habitacionId, huespedId."),
        ("CONSUMOS_EXTRAS", "Cargos adicionales.", "categoria, descripcion, monto, estado, reservaSnapshot."),
        ("PAGOS", "Movimientos confirmados.", "monto, método, estado, reservaSnapshot."),
        ("SOLICITUDES_RECUPERACION_CLAVE", "Auditoría de recuperación.", "userId, email, status, requestedAt, resolvedBy."),
    ], widths=[4, 5, 7], caption="Tabla 12. Colecciones de Firestore y campos representativos.")
    add_title(doc, "Anexo C. Procedimiento de reproducción con seed")
    add_text(doc, [
        "1. Crear o seleccionar un proyecto Firebase de desarrollo con Cloud Firestore habilitado. 2. Configurar firebase-credentials.json local o una variable de cuenta de servicio. 3. Instalar dependencias con npm ci. 4. Definir SEED_ADMIN_PASSWORD y SEED_RECEPTIONIST_PASSWORD. 5. Ejecutar npm run db:seed -- --apply. 6. Iniciar el servidor y acceder con las cuentas de demostración. 7. Cambiar las claves en el primer inicio de sesión.",
        "El procedimiento no debe ejecutarse sobre una base de producción sin respaldo y autorización. El seed crea sólo identificadores demo- y omite documentos existentes, por lo que no tiene una función de borrado ni de restauración de datos reales.",
    ])
    add_title(doc, "Anexo D. Lista de verificación pre-final")
    add_table(doc, ["Verificación", "Resultado a registrar", "Responsable"], [
        ("Repositorio actualizado", "Rama main, commit y enlace comprobables.", "Autor."),
        ("Render disponible", "Landing, /app/login y /api/health responden.", "Autor."),
        ("Usuarios de demostración", "ADMIN y RECEPCIONISTA pueden iniciar sesión.", "Autor."),
        ("Flujo de reserva", "Alta, solapamiento, check-in y check-out verificados.", "Autor."),
        ("Flujo económico", "Consumo, cuenta y pago demostrados.", "Autor."),
        ("Tesis", "Mínimo 50 páginas, índices actualizados y APA revisada.", "Autor."),
        ("Presentación y video", "PPTX terminado y grabación ensayada con cámara.", "Autor."),
    ], widths=[4, 8, 4], caption="Tabla 13. Checklist antes de la instancia pre-final.")
    add_title(doc, "Anexo E. Guion sugerido para la exposición")
    add_table(doc, ["Tiempo estimado", "Contenido", "Evidencia visual"], [
        ("0–3 min", "Problema, objetivo y alcance.", "Portada y arquitectura."),
        ("3–7 min", "Tecnologías, roles y seguridad.", "Diagrama y matriz de permisos."),
        ("7–17 min", "Demostración: login, reserva, consumo y pago.", "Aplicación en Render."),
        ("17–21 min", "Recuperación interna, seed y despliegue.", "Módulo Usuarios, README y health check."),
        ("21–25 min", "Resultados, límites, conclusiones y futuro.", "Tesis y diapositiva final."),
    ], widths=[3, 8, 5], caption="Tabla 14. Secuencia orientativa para video o defensa oral.")
    add_title(doc, "Anexo F. Registro de evidencia de pruebas")
    add_text(doc, [
        "La siguiente matriz se utiliza para registrar la evidencia realizada en el entorno local y en el despliegue. Antes del pre-final se debe completar una fila por cada caso, adjuntar una captura cuando corresponda y consignar si el resultado fue satisfactorio o si queda una incidencia abierta.",
    ])
    add_table(doc, ["Fecha", "Caso", "Entorno", "Resultado", "Evidencia"], [
        ("____/____/2026", "PF-01", "Local / Render", "Pendiente de completar.", "Captura o URL."),
        ("____/____/2026", "PF-02", "Local / Render", "Pendiente de completar.", "Captura o URL."),
        ("____/____/2026", "PF-03", "Local / Render", "Pendiente de completar.", "Captura o URL."),
        ("____/____/2026", "PF-04", "Local / Render", "Pendiente de completar.", "Captura o URL."),
        ("____/____/2026", "PF-05", "Local / Render", "Pendiente de completar.", "Captura o URL."),
        ("____/____/2026", "PF-06", "Local / Render", "Pendiente de completar.", "Captura o URL."),
    ], widths=[2.7, 2, 3.3, 4, 4], caption="Tabla 15. Plantilla de evidencia para los casos de prueba.")
    add_title(doc, "Anexo G. Contrato resumido de API")
    add_text(doc, [
        "La API se organiza bajo el prefijo /api. Los endpoints protegidos requieren Authorization: Bearer <token>. Los cuerpos se intercambian en JSON y los errores se expresan mediante códigos HTTP y un mensaje para el cliente. Esta tabla no reemplaza una especificación OpenAPI, pero permite demostrar el contrato principal durante la evaluación.",
    ])
    add_table(doc, ["Método y ruta", "Propósito", "Autorización"], [
        ("POST /api/auth/login", "Inicia sesión y entrega JWT.", "Público."),
        ("POST /api/auth/forgot-password", "Solicita revisión interna de acceso.", "Público, respuesta genérica."),
        ("GET /api/reservas", "Lista reservas con estado operativo.", "ADMIN o RECEPCIONISTA."),
        ("POST /api/consumos", "Registra un consumo pendiente.", "ADMIN o RECEPCIONISTA."),
        ("PATCH /api/consumos/:id/incluir-en-cuenta", "Incluye el cargo en saldo.", "ADMIN o RECEPCIONISTA."),
        ("POST /api/pagos", "Registra un pago confirmado.", "ADMIN o RECEPCIONISTA."),
        ("PATCH /api/users/:id/temporary-password", "Asigna clave temporal.", "ADMIN."),
    ], widths=[5.5, 6.5, 4], caption="Tabla 16. Endpoints representativos del contrato de API.")
    add_title(doc, "Anexo H. Configuración segura por entorno")
    add_text(doc, [
        "La configuración se separa del código fuente mediante archivos .env locales y variables del panel de Render. El repositorio incluye ejemplos con valores vacíos; nunca se debe usar un secreto real como ejemplo ni adjuntarlo a la tesis. La cuenta de servicio de Firebase permite que el backend acceda a Firestore y debe mantenerse exclusivamente del lado servidor.",
    ])
    add_table(doc, ["Variable", "Entorno local", "Render", "Regla de seguridad"], [
        ("JWT_SECRET", ".env", "Variable secreta", "Clave aleatoria, no versionada."),
        ("FIREBASE_SERVICE_ACCOUNT_BASE64", "Opcional", "Variable secreta", "No imprimir ni compartir."),
        ("FIREBASE_SERVICE_ACCOUNT_JSON", "Opcional", "Alternativa", "No guardar en repositorio."),
        ("CORS_ORIGINS", ".env", "Variable de configuración", "Declarar sólo orígenes requeridos."),
        ("VITE_API_URL", "DVT-SoftwareHotelero/.env", "No necesaria si comparte origen", "No contiene secreto."),
    ], widths=[4.2, 3.4, 3.6, 4.8], caption="Tabla 17. Variables de entorno y reglas de custodia.")
    add_title(doc, "Anexo I. Plan de contingencia para la demostración")
    add_text(doc, [
        "El plan de contingencia no reemplaza las pruebas, pero reduce la probabilidad de que una situación externa opaque la demostración. La instancia gratuita de Render se debe abrir antes de comenzar para permitir el arranque si estuvo inactiva. También se debe conservar una copia local del repositorio con dependencias instaladas y un conjunto de datos ficticios verificable.",
    ])
    add_table(doc, ["Situación", "Acción inmediata", "Alternativa"], [
        ("Render tarda en responder", "Abrir /api/health y esperar el inicio de instancia.", "Continuar con explicación de arquitectura y luego retomar demo."),
        ("Datos de prueba inconsistentes", "Usar proyecto Firebase de demostración y seed demo-.", "Mostrar procedimiento de réplica documentado."),
        ("Falla de internet", "Mantener aplicación local preparada.", "Demostrar build local y grabación breve de respaldo."),
        ("Olvido de credenciales", "Usar cuenta de prueba autorizada.", "Demostrar flujo de clave temporal con ADMIN."),
        ("Cambio sin desplegar", "Comparar commit de main con Render.", "No presentar hasta verificar build y health check."),
    ], widths=[4, 6, 6], caption="Tabla 18. Plan de contingencia para la exposición.")


def build_document():
    figures = build_figures()
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_body(doc, figures)
    add_references_and_annexes(doc)
    doc.core_properties.title = "DVT — Software de Gestión Hotelera"
    doc.core_properties.subject = "Proyecto Final Integrador"
    doc.core_properties.author = "Iván Federico Ibarra"
    doc.core_properties.comments = "Borrador técnico para instancia pre-final."
    doc.save(OUTPUT)
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    build_document()
